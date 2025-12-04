from pathlib import Path
from typing import List # 用于类型注解
from langchain_core.documents import Document # 从 LangChain 核心库导入 Document 类。
# Document 是 LangChain 中表示“文本内容 + 元数据”的标准结构
from langchain_text_splitters import RecursiveCharacterTextSplitter
# 导入 LangChain 的文本分割器，用于将长文档切分成较小的块（chunks），便于后续向量化或检索。
from pypdf import PdfReader # 用于读取PDF文件内容
import docx
from app.config import settings

def load_pdf(path: Path) -> List[Document]:
    reader = PdfReader(str(path))
    docs = []
    for i, page in enumerate(reader.pages): # 便利PDF的每一页(reader.pages是页面迭代器), i是页码索引
        text = page.extract_text() or "" # 从当前页提取文本
        if text.strip(): # 提取的文本去除空白后非空(避免添加空页)
            docs.append(Document( # 创建一个Document对象
                page_content=text, # page_content:提取的文本
                metadata={"source": str(path), "page": i+1} # 包含文件路径和页码(页码从1开始为i+1)
            ))
    return docs

def load_docx(path: Path) -> List[Document]:
    d = docx.Document(str(path))
    # 遍历文档中所有段落,过滤掉空段落(p.text.strip()非空),用换行符PN结所有段落文本
    text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
    # 如果 text 非空，返回包含一个 Document 的列表,否则返回空列表（避免加载空文档）
    return [Document(page_content=text, metadata={"source": str(path)})] if text else []

def load_docs(dir_path: str) -> List[Document]:
    p = Path(dir_path)
    docs: List[Document] = []
    for f in p.rglob("*"): # 使用 rglob("*") 递归遍历目录下所有文件和子目录中的文件
        if f.suffix.lower() == ".pdf":
            docs.extend(load_pdf(f))
            # 如果是 PDF 文件，调用 load_pdf，并用 extend 将返回的多个 Document 加入总列表
        elif f.suffix.lower() in [".docx", ".doc"]:
            docs.extend(load_docx(f))
        elif f.suffix.lower() in [".md", ".txt"]:
            docs.append(Document(page_content=f.read_text(encoding="utf-8"),
                                 metadata={"source": str(f)}))
    return docs

def split_docs(docs: List[Document]) -> List[Document]:
    """
    定义文本分割函数，输入文档列表，输出分割后的更小文档列表
    """
    splitter = RecursiveCharacterTextSplitter(
        # 创建一个递归字符文本分割器
        chunk_size=settings.chunk_size, # 每个块的最大字符数
        chunk_overlap=settings.chunk_overlap # 相邻块之间的重叠字符数
    )
    # 调用 split_documents 方法，自动按规则切分每个 Document 的 page_content
    return splitter.split_documents(docs)

def load_single_file(path: Path) -> List[Document]:
    """根据文件后缀加载文件，返回LangChain的 Document列表"""
    suf = path.suffix.lower()
    if suf == ".pdf":
        return load_pdf(path)
    if suf in [".docx", ".doc"]:
        return load_docx(path)
    if suf in [".md", ".txt"]:
        text = path.read_text(encoding="utf-8")
        return [Document(page_content=text, metadata={"source": str(path)})] if text.strip() else []
    return []

# 第二个函数主要是把一批Document切成小块，并且给每一小块贴上权限标签visibility和文档ID。
def split_with_visibility(docs: List[Document], visibility: str, doc_id: str | None = None) -> List[Document]:
    chunks = split_docs(docs)
    for c in chunks:
        c.metadata = dict(c.metadata or {})
        c.metadata["visibility"] = visibility
        if doc_id:
            c.metadata["doc_id"] = doc_id
    return chunks

if __name__ == "__main__":
    docs = split_docs(load_docs("/home/mrwu/PycharmProjects/enterprise-kb-assistant/data/docs"))
    for _ in docs:
        print(_)
